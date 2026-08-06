"""
简历文件解析 — 提取文本 + LLM 转结构化 JSON

支持: PDF, Word (.docx), Markdown, 纯文本, JSON
"""
from __future__ import annotations

import io
import json
import logging
from pathlib import Path

import httpx
from openai import OpenAI

from config import settings

logger = logging.getLogger(__name__)

client = OpenAI(
    api_key=settings.llm_api_key,
    base_url=settings.llm_base_url,
    http_client=httpx.Client(transport=httpx.HTTPTransport(retries=0), timeout=httpx.Timeout(90.0, connect=10.0)),
)

PARSE_PROMPT = """你是简历解析专家。把以下从简历文件中提取的原始文本，转换为标准的结构化简历 JSON。

## 输出 JSON 结构

```json
{
  "name": "姓名",
  "title": "当前/目标职位",
  "years": 5,
  "summary": "1-2句个人总结",
  "skills": ["技能1", "技能2"],
  "target": {
    "roles": ["目标岗位"],
    "salary_min": 15000,
    "salary_max": 25000,
    "locations": ["期望城市"],
    "industries": ["目标行业"],
    "remote_only": false
  },
  "experience": [
    {
      "company": "公司名",
      "title": "职位",
      "start": "2020-03",
      "end": "2023-06",
      "highlights": ["亮点1", "亮点2"]
    }
  ],
  "education": [
    {
      "school": "学校名",
      "degree": "本科/硕士/博士",
      "major": "专业",
      "year": 2019
    }
  ],
  "projects": [],
  "certifications": []
}
```

## 规则
- 从文本中提取真实信息，绝不编造
- years 根据工作经验年限推算
- 文本中没提到的字段用空值（null/[]/""）
- salary_min/max 根据文本中提到的期望薪资换算为月薪数字（元），没提到就给 null
- 只输出 JSON，不要其他文字"""

# 支持的文件格式
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown", ".json"}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def extract_text(filename: str, file_bytes: bytes) -> str:
    """根据文件扩展名提取文本内容"""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_bytes)
    elif ext in (".txt", ".md", ".markdown"):
        return _extract_text(file_bytes)
    elif ext == ".json":
        return _extract_text(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    result = "\n".join(text_parts)
    if not result.strip():
        raise ValueError("PDF 无法提取文本内容，可能是扫描件或图片 PDF")
    return result


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _extract_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def parse_to_markdown(filename: str, file_bytes: bytes) -> str:
    """完整流程：提取文本 → LLM 整理为 Markdown"""
    text = extract_text(filename, file_bytes)

    if not text.strip():
        raise ValueError("无法从文件中提取内容，文件可能为空")

    # 如果已经是 Markdown 文件，直接返回
    ext = Path(filename).suffix.lower()
    if ext in (".md", ".markdown"):
        return text

    # 如果是纯文本但接近 Markdown 格式，直接返回
    if ext == ".txt" and ("#" in text or "##" in text):
        return text

    # LLM 转换为 Markdown
    logger.info("Parsing resume to Markdown with LLM, text length: %d", len(text))
    text = text[:12000]

    response = client.chat.completions.create(
        model=settings.llm_model,
        messages=[
            {"role": "system", "content": """你是简历格式整理助手。把以下从简历文件中提取的原始文本，整理为规范的 Markdown 格式。

## 输出格式

```
# 姓名 — 目标岗位

个人总结（1-2句，如果没有则跳过）

## 技能

- 熟练掌握: 技能1 / 技能2
- 熟悉: 技能3
- 了解: 技能4

## 工作经历

### 公司名 — 职位（时间段）

- 亮点1（量化结果）
- 亮点2

## 项目经历

### 项目名 — 角色（时间段）

- 亮点1
- 亮点2

## 教育背景

- **学校名** — 学位 专业（年份）

## 证书

- 证书1
- 证书2
```

## 规则
1. 只整理格式，不改变信息内容，不编造
2. 原文中没提到的部分直接跳过不输出
3. 只输出整理后的 Markdown，不要任何解释"""},
            {"role": "user", "content": f"## 简历原始文本\n\n{text}\n\n请整理为规范 Markdown。"},
        ],
        temperature=0.2,
    )

    return (response.choices[0].message.content or "").strip()


def parse_to_json(filename: str, file_bytes: bytes) -> dict:
    """完整流程：提取文本 → LLM 转 JSON"""
    # Step 1: 提取文本
    text = extract_text(filename, file_bytes)

    if not text.strip():
        raise ValueError("无法从文件中提取内容，文件可能为空")

    # Step 2: 如果已经是 JSON，直接解析
    ext = Path(filename).suffix.lower()
    if ext == ".json":
        try:
            data = json.loads(text)
            # 确保有基本字段
            data.setdefault("name", "")
            data.setdefault("skills", [])
            data.setdefault("experience", [])
            data.setdefault("education", [])
            return data
        except json.JSONDecodeError:
            pass  # JSON 解析失败，走 LLM 流程

    # Step 3: LLM 转换
    logger.info("Parsing resume with LLM, text length: %d", len(text))
    text = text[:12000]  # 限制输入长度

    response = client.chat.completions.create(
        model=settings.llm_model,
        messages=[
            {"role": "system", "content": PARSE_PROMPT},
            {"role": "user", "content": f"## 简历原始文本\n\n{text}\n\n请转换为结构化 JSON。"},
        ],
        temperature=0.2,
    )

    raw = response.choices[0].message.content or ""
    json_str = _extract_json(raw)

    try:
        data = json.loads(json_str)
    except json.JSONDecodeError as e:
        logger.error("Failed to parse LLM output as JSON: %s", raw[:500])
        raise ValueError("AI 解析失败，请手动粘贴简历内容") from e

    # 确保必要字段存在
    data.setdefault("name", "")
    data.setdefault("title", "")
    data.setdefault("skills", [])
    data.setdefault("experience", [])
    data.setdefault("education", [])

    return data


def _extract_json(raw: str) -> str:
    """从 LLM 响应中提取 JSON"""
    raw = raw.strip()
    if raw.startswith("```"):
        idx = raw.find("\n")
        if idx != -1:
            raw = raw[idx + 1:]
        if raw.rstrip().endswith("```"):
            raw = raw.rstrip()[:raw.rstrip().rfind("```")]
    start = raw.find("{")
    end = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        raw = raw[start:end + 1]
    return raw.strip()
